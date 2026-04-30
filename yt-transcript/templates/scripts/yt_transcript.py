#!/usr/bin/env python3
"""
yt-transcript — YouTube videosunu okunabilir dokümana dönüştürür.
Kullanım: python yt_transcript.py <URL> [--format word|pdf] [--lang tr] [--output ./] [--whisper-model base]
"""

import argparse
import os
import re
import sys
import tempfile
import textwrap
from datetime import datetime
from pathlib import Path


def extract_video_id(url: str) -> str:
    patterns = [
        r"(?:v=|youtu\.be/|embed/|shorts/)([a-zA-Z0-9_-]{11})",
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    raise ValueError(f"Geçerli YouTube URL'i bulunamadı: {url}")


def get_transcript_from_api(video_id: str, lang: str) -> tuple[str, str]:
    """youtube-transcript-api ile altyazı çeker. (metin, başlık) döner."""
    from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptFound, TranscriptsDisabled

    api = YouTubeTranscriptApi()

    # Önce istenen dil, sonra İngilizce, sonra her ne varsa
    lang_priority = [lang, "en", "tr"]
    transcript_list = api.list(video_id)

    transcript = None
    for l in lang_priority:
        try:
            transcript = transcript_list.find_transcript([l])
            break
        except Exception:
            continue

    if transcript is None:
        # İlk mevcut olanı al
        try:
            available = list(transcript_list)
            if not available:
                raise NoTranscriptFound(video_id, [], {})
            transcript = available[0]
        except Exception as e:
            raise NoTranscriptFound(video_id, [], {}) from e

    fetched = transcript.fetch()
    # FetchedTranscript objesi — snippet listesi
    segments = fetched if isinstance(fetched, list) else list(fetched)
    raw_text = " ".join(s.get("text", "") if isinstance(s, dict) else s.text for s in segments)
    return raw_text, ""


def get_transcript_via_whisper(url: str, model_size: str) -> str:
    """yt-dlp + whisper ile transkripsiyon yapar."""
    import subprocess
    import whisper

    print(f"  Ses indiriliyor: {url}", flush=True)
    with tempfile.TemporaryDirectory() as tmpdir:
        audio_path = os.path.join(tmpdir, "audio.mp3")
        result = subprocess.run(
            ["yt-dlp", "-x", "--audio-format", "mp3", "-o", audio_path, url],
            capture_output=True, text=True
        )
        if result.returncode != 0:
            raise RuntimeError(f"yt-dlp hatası: {result.stderr}")

        print(f"  Whisper '{model_size}' modeli ile transkripsiyon yapılıyor...", flush=True)
        model = whisper.load_model(model_size)
        output = model.transcribe(audio_path)
        return output["text"]


def get_video_title(video_id: str) -> str:
    try:
        import subprocess
        result = subprocess.run(
            ["yt-dlp", "--get-title", f"https://www.youtube.com/watch?v={video_id}"],
            capture_output=True, text=True, timeout=15
        )
        if result.returncode == 0:
            return result.stdout.strip()
    except Exception:
        pass
    return f"YouTube Video ({video_id})"


CHUNK_SIZE = 12000  # karakter (~3,000 token) — çıktı için bol yer bırakır


def _call_openrouter(client, model: str, api_key: str, system_prompt: str, user_prompt: str) -> str:
    response = client.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://github.com/sultanxgokce/Sx-Claude-Skills",
        },
        json={
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        },
    )
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]


def format_with_claude(raw_text: str, title: str, url: str) -> str:
    """OpenRouter üzerinden Claude ile metni chunk'lara bölerek formatlar."""
    import httpx

    api_key = os.environ.get("OPENROUTER_API_KEY", "")
    if not api_key:
        print("  UYARI: OPENROUTER_API_KEY bulunamadı. Ham metin kullanılıyor.", flush=True)
        return raw_text

    model = os.environ.get("OPENROUTER_MODEL", "anthropic/claude-haiku-4-5")

    system_prompt = textwrap.dedent("""
        Sen bir içerik editörüsün. Sana YouTube video transkriptinin bir bölümü verilecek.
        Görevin: ham, bitimsiz transkripti okunabilir, düzenli bir makaleye dönüştürmek.

        Kurallar:
        - Konuşma dilini koru, metni robotik hale getirme
        - Mantıklı paragraflara böl (her paragraf bir fikir)
        - Ana konular için ## başlık kullan
        - Tekrarları temizle ama içeriği KISALTMA — her fikri koru
        - "eee", "şey", "yani" gibi dolgu kelimeleri kaldır
        - Türkçe ise Türkçe kalsın, İngilizce ise İngilizce kalsın
        - Sadece düzenlenmiş metni döndür, yorum ekleme
    """).strip()

    # Metni CHUNK_SIZE karakterlik parçalara böl (kelime sınırında)
    words = raw_text.split()
    chunks = []
    current = []
    current_len = 0
    for word in words:
        current.append(word)
        current_len += len(word) + 1
        if current_len >= CHUNK_SIZE:
            chunks.append(" ".join(current))
            current = []
            current_len = 0
    if current:
        chunks.append(" ".join(current))

    total = len(chunks)
    print(f"  Transkript {total} parçaya bölündü ({len(raw_text):,} karakter).", flush=True)

    formatted_parts = []
    with httpx.Client(timeout=180) as client:
        for i, chunk in enumerate(chunks, 1):
            print(f"  Parça {i}/{total} formatlanıyor...", flush=True)
            user_prompt = f"Video başlığı: {title}\nBölüm {i}/{total}\n\nTranskript:\n{chunk}"
            part = _call_openrouter(client, model, api_key, system_prompt, user_prompt)
            formatted_parts.append(part)

    return "\n\n".join(formatted_parts)


def save_as_word(formatted_text: str, title: str, url: str, output_path: Path):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Sayfa kenar boşlukları
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Başlık
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Meta bilgi
    meta = doc.add_paragraph()
    meta.add_run(f"Kaynak: ").bold = True
    meta.add_run(url)
    meta.add_run(f"\nOluşturulma: ").bold = True
    meta.add_run(datetime.now().strftime("%d.%m.%Y %H:%M"))
    doc.add_paragraph()

    # İçerik
    for line in formatted_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("---"):
            doc.add_paragraph("─" * 40)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

    doc.save(str(output_path))


def save_as_pdf(formatted_text: str, title: str, url: str, output_path: Path):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(25, 20, 20)

    # Başlık
    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(0, 10, title, align="L")
    pdf.ln(3)

    # Meta
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 5, f"Kaynak: {url}", ln=True)
    pdf.cell(0, 5, f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')}", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(5)
    pdf.line(25, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(5)

    # İçerik
    for line in formatted_text.split("\n"):
        line = line.strip()
        if not line:
            pdf.ln(3)
            continue
        if line.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 7, line[3:], align="L")
            pdf.ln(2)
        elif line.startswith("# "):
            pdf.set_font("Helvetica", "B", 15)
            pdf.multi_cell(0, 8, line[2:], align="L")
            pdf.ln(2)
        elif line.startswith("---"):
            pdf.line(25, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(4)
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, line, align="L")
            pdf.ln(1)

    pdf.output(str(output_path))


def main():
    parser = argparse.ArgumentParser(description="YouTube videosunu dokümana dönüştür")
    parser.add_argument("url", help="YouTube video URL'i")
    parser.add_argument("--format", choices=["word", "pdf"], default="word", help="Çıktı formatı (varsayılan: word)")
    parser.add_argument("--lang", default="tr", help="Altyazı dil kodu (varsayılan: tr)")
    parser.add_argument("--output", default="./", help="Çıktı dizini")
    parser.add_argument("--whisper-model", default="base", choices=["tiny", "base", "small", "medium", "large"])
    args = parser.parse_args()

    print(f"\nYouTube Transcript — {args.url}\n{'─'*50}")

    video_id = extract_video_id(args.url)
    print(f"  Video ID: {video_id}")

    # Başlık al
    print("  Video başlığı alınıyor...")
    title = get_video_title(video_id)
    print(f"  Başlık: {title}")

    # Transkript al
    raw_text = None
    print("  Altyazı aranıyor...")
    try:
        raw_text, _ = get_transcript_from_api(video_id, args.lang)
        print("  Altyazı bulundu.")
    except Exception as e:
        print(f"  Altyazı bulunamadı ({e}), Whisper devreye alınıyor...")
        try:
            raw_text = get_transcript_via_whisper(args.url, args.whisper_model)
        except Exception as e2:
            print(f"  Whisper de başarısız: {e2}")
            sys.exit(1)

    # Claude ile formatla
    formatted = format_with_claude(raw_text, title, args.url)

    # Dosya adı
    safe_title = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "_")[:50]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    ext = "docx" if args.format == "word" else "pdf"
    filename = f"{safe_title}_{timestamp}.{ext}"
    output_path = Path(args.output) / filename

    # Kaydet
    print(f"  {args.format.upper()} olarak kaydediliyor: {output_path}")
    if args.format == "word":
        save_as_word(formatted, title, args.url, output_path)
    else:
        save_as_pdf(formatted, title, args.url, output_path)

    print(f"\nTamamlandı: {output_path.resolve()}\n")


if __name__ == "__main__":
    main()
